"""
FRAS Project Documentation Generator
Generates a comprehensive Word (.docx) document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────── helpers ───────────────

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def heading(doc, text, level, color_hex="0a2342"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    h.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    h.paragraph_format.space_after  = Pt(4)
    return h


def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Pt(36 + level * 18)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Pt(36)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    # light blue background paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EFF6FF')
    pPr.append(shd)
    return p


def info_table(doc, rows):
    """Two-column label/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)
    for i, (label, value) in enumerate(rows):
        lc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        set_cell_bg(lc, 'DBEAFE')
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        vp = vc.paragraphs[0]
        vp.add_run(value).font.size = Pt(10)
    doc.add_paragraph()


def section_table(doc, headers, data_rows, header_bg='1D4ED8'):
    cols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        rp = cell.paragraphs[0]
        run = rp.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for ri, row_data in enumerate(data_rows):
        row = table.rows[ri + 1]
        bg = 'F0F9FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].add_run(val).font.size = Pt(9.5)
    doc.add_paragraph()

# ═══════════════════════════════════════
#   BUILD DOCUMENT
# ═══════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover Page ──────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(80)
tr = title_p.add_run('FRAS – Face Recognition Attendance System')
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = RGBColor(0x0a, 0x23, 0x42)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('Complete Technical & Functional Documentation')
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x00, 0xb8, 0xa9)

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta_p.add_run(
    f"Version 1.1.4  |  Platform: Android & iOS  |  Date: {datetime.date.today().strftime('%B %d, %Y')}"
)
mr.font.size = Pt(11)
mr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_page_break()

# ── Table of Contents placeholder ──────────────────────────────────────────────
heading(doc, 'Table of Contents', 1)
toc_items = [
    ('1', 'Project Overview'),
    ('2', 'Project Architecture'),
    ('3', 'Technology Stack & Libraries'),
    ('4', 'Database Schema & Data Models'),
    ('5', 'Core Algorithms & Techniques'),
    ('6', 'Module-by-Module Breakdown'),
    ('7', 'Screen / UI Documentation'),
    ('8', 'Navigation & Application Flow'),
    ('9', 'API Integration'),
    ('10', 'Authentication & Security'),
    ('11', 'State Management'),
    ('12', 'Settings & Configuration'),
    ('13', 'Performance Optimisations'),
    ('14', 'Background Services'),
    ('15', 'Known Issues & Limitations'),
    ('16', 'Build, Run & Deployment'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'  {num}.  {title}')
    r.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '1. Project Overview', 1)
add_horizontal_rule(doc)

body(doc,
    'FRAS (Face Recognition Attendance System) is a production-grade React Native mobile '
    'application designed to automate employee attendance tracking using biometric face '
    'recognition. The system operates offline-first, stores all data locally in an SQLite '
    'database, and synchronises with a cloud backend when an internet connection is available.',
)

info_table(doc, [
    ('Project Name',    'FRAS – Face Recognition Attendance System'),
    ('Version',         '1.1.4'),
    ('Framework',       'React Native 0.80.0 + Expo 53'),
    ('React Version',   '19.1.0'),
    ('Platforms',       'Android (primary), iOS (supported)'),
    ('Architecture',    'Offline-first, Service-oriented'),
    ('State Mgmt',      'React Context API + AsyncStorage + In-memory cache'),
    ('Local Storage',   'SQLite (react-native-sqlite-storage)'),
    ('Cloud Backend',   'REST API (HTTP/JSON)'),
    ('AI / ML',         'TensorFlow Lite (native), Google ML Kit Face Detection'),
])

# ════════════════════════════════════════════════════════════════════════════════
#  2. PROJECT ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '2. Project Architecture', 1)
add_horizontal_rule(doc)

body(doc, 'The application is organized into clearly separated layers:', bold=True)

arch_data = [
    ('Layer',           'Folder / Files',                       'Responsibility'),
    ('Screens (UI)',    'screens/',                             'Navigation targets; user-facing views'),
    ('Components',      'components/',                          'Reusable UI widgets and modals'),
    ('Services',        'services/',                            'Business logic (recognition, sync, auth)'),
    ('Repositories',    'database/*.repository.js',             'SQLite CRUD operations per domain'),
    ('Database',        'database/connection.js, schema.js',   'SQLite setup, migrations, indexes'),
    ('Utils',           'utils/',                               'Algorithms, GPS, settings helpers'),
    ('Config',          'config/config.ts, constants/',         'API base URL, colour palette'),
    ('Context',         'AuthContext.js',                       'Global authentication state'),
    ('Assets',          'assets/',                              'Images, audio files'),
]
section_table(doc, ['Layer', 'Folder / Files', 'Responsibility'], arch_data[1:])

heading(doc, '2.1 Folder Structure', 2)
code_block(doc,
"""FRAS-dev/
├── assets/
│   ├── images/          # App logo, UI artwork
│   └── sounds/          # success.mp3, warning.mp3
├── components/          # Reusable React Native components
│   ├── AttendanceComps/ # Attendance-specific popups
│   ├── FacePositionOverlay/
│   ├── FaceRecognition/
│   └── GridImageCapture/
├── config/config.ts     # API base URL
├── constants/colors.ts  # Colour tokens
├── database/            # SQLite layer
├── module/resources/    # Resource management sub-module
├── screens/             # App screens
├── services/            # Business logic services
├── utils/               # Algorithm utilities
├── App.js               # Root component
├── AuthContext.js        # Auth context provider
└── index.js             # Entry point""")

# ════════════════════════════════════════════════════════════════════════════════
#  3. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '3. Technology Stack & Libraries', 1)
add_horizontal_rule(doc)

heading(doc, '3.1 Core Framework', 2)
info_table(doc, [
    ('React Native', '0.80.0 – cross-platform mobile framework'),
    ('React',        '19.1.0 – UI rendering library'),
    ('Expo',         '53.0.13 – managed workflow tooling'),
    ('TypeScript',   'Optional – tsconfig.json present; JS files dominant'),
])

heading(doc, '3.2 Navigation', 2)
section_table(doc,
    ['Package', 'Version', 'Purpose'],
    [
        ('@react-navigation/native',      '^7.1.14', 'Navigation container & core'),
        ('@react-navigation/stack',       '^7.4.2',  'Stack (sequential) navigation'),
        ('@react-navigation/bottom-tabs', '^7.4.2',  'Bottom tab bar'),
        ('react-native-screens',          '^4.11.1', 'Native screen optimisation'),
        ('react-native-gesture-handler',  '^2.26.0', 'Gesture support'),
        ('react-native-safe-area-context','^^5.5.0', 'Safe area insets'),
    ]
)

heading(doc, '3.3 Camera & Computer Vision', 2)
section_table(doc,
    ['Package', 'Version', 'Purpose'],
    [
        ('react-native-vision-camera',         '^4.7.0',  'High-performance camera access'),
        ('@react-native-ml-kit/face-detection', '^1.3.2',  'Real-time face detection (Google ML Kit)'),
        ('react-native-image-resizer',          '^1.4.5',  'Image resizing before embedding extraction'),
        ('@react-native-community/image-editor','^^4.3.0', 'Face crop from camera frame'),
    ]
)

heading(doc, '3.4 Storage & Data', 2)
section_table(doc,
    ['Package', 'Version', 'Purpose'],
    [
        ('@react-native-async-storage/async-storage', '^2.2.0', 'Persistent key-value store'),
        ('react-native-sqlite-storage',               '^6.0.1',  'Local relational database'),
        ('react-native-fs',                           '^2.20.0', 'File system access'),
        ('uuid',                                      '^8.3.2',  'UUID generation'),
    ]
)

heading(doc, '3.5 Machine Learning', 2)
section_table(doc,
    ['Package / Module', 'Version', 'Purpose'],
    [
        ('TFLiteFaceModule (native)',    'Custom',  '512-D face embedding extraction (TF Lite)'),
        ('TFLiteModule (native)',        'Custom',  'General TensorFlow Lite inference'),
        ('TSFModuleIOS (native)',        'Custom',  'iOS-specific TF Lite wrapper'),
        ('FaceSpoofDetector (native)',   'Custom',  'Android anti-spoofing detection'),
        ('@react-native-ml-kit/face-detection', '^1.3.2', 'Face bounding box & landmarks'),
        ('hnsw',                         '^1.1.1', 'HNSW ANN index (in development)'),
    ]
)

heading(doc, '3.6 UI & Animation', 2)
section_table(doc,
    ['Package', 'Version', 'Purpose'],
    [
        ('lucide-react-native',           '^0.525.0', 'Icon set'),
        ('react-native-vector-icons',     '^10.2.0',  'Additional icon set'),
        ('react-native-svg',              '^15.12.0', 'SVG rendering'),
        ('react-native-circular-progress','^^1.4.1',  'Circular progress indicator'),
        ('react-native-reanimated',       '^3.18.0',  'Native-thread animations'),
        ('react-native-worklets-core',    '^1.6.0',   'JS worklets for VisionCamera'),
    ]
)

heading(doc, '3.7 Location & Network', 2)
section_table(doc,
    ['Package', 'Version', 'Purpose'],
    [
        ('react-native-geolocation-service', '^5.3.1', 'GPS location tracking'),
        ('@react-native-community/netinfo',  '^11.4.1', 'Network connectivity monitoring'),
    ]
)

heading(doc, '3.8 Device & Logging', 2)
section_table(doc,
    ['Package', 'Version', 'Purpose'],
    [
        ('react-native-device-info',         '^14.1.1', 'Device metadata (OS, model)'),
        ('react-native-orientation-locker',  '^1.5.0',  'Lock screen orientation'),
        ('@sayem314/react-native-keep-awake','^^1.4.0', 'Prevent screen sleep during recognition'),
        ('react-native-sound-player',        '^0.14.5', 'Audio feedback (success/warning)'),
        ('react-native-background-fetch',    '^4.3.0',  'Background sync tasks'),
        ('@bugfender/rn-bugfender',          '^5.0.1',  'Remote error/crash logging'),
        ('vexo-analytics',                   '^1.5.3',  'App usage analytics'),
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  4. DATABASE SCHEMA
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '4. Database Schema & Data Models', 1)
add_horizontal_rule(doc)

body(doc,
    'The app uses SQLite (via react-native-sqlite-storage) as its local relational database. '
    'WAL mode and foreign keys are enabled. Four main tables handle face vectors, attendance '
    'records, sync queues, and user preferences.'
)

# --- FACEVECTOR ---
heading(doc, '4.1 facevector – Enrolled Face Templates', 2)
section_table(doc,
    ['Column', 'Type', 'Description'],
    [
        ('id',         'INTEGER PK AUTOINCREMENT', 'Internal row ID'),
        ('uuid',       'TEXT UNIQUE NOT NULL',     'Employee GUID (primary identifier)'),
        ('staffid',    'TEXT',                     'Employee staff / payroll ID'),
        ('name',       'TEXT',                     'Employee display name'),
        ('vector',     'TEXT',                     'Average 512-D embedding (JSON array)'),
        ('vectors',    'TEXT',                     'Up to 5 raw embeddings (JSON array of arrays)'),
        ('img',        'TEXT',                     'Base64-compressed face photo'),
        ('enrollmode', 'TEXT DEFAULT online',      '"online" | "offline" enrollment mode'),
        ('createdby',  'TEXT',                     'UUID of the admin who enrolled'),
        ('syncdate',   'DATETIME',                 'Timestamp of last cloud sync'),
    ]
)

# --- PUNCHRECORD ---
heading(doc, '4.2 punchrecord – Attendance Records', 2)
section_table(doc,
    ['Column', 'Type', 'Description'],
    [
        ('id',              'TEXT PK',          'UUID record identifier'),
        ('uuid',            'TEXT NOT NULL',    'Employee GUID (FK → facevector.uuid)'),
        ('projectid',       'TEXT',             'Assigned project / site'),
        ('punchtype',       'TEXT',             '"in" or "out"'),
        ('punchdate',       'DATETIME NOT NULL','Local device timestamp'),
        ('lat',             'TEXT',             'GPS latitude'),
        ('lan',             'TEXT',             'GPS longitude (note: column named "lan")'),
        ('attendancetype',  'TEXT',             'Role / department category'),
        ('punchmode',       'TEXT DEFAULT offline', '"offline" | "synced"'),
        ('syncdate',        'DATETIME',         'Timestamp when pushed to server'),
        ('retry_count',     'INTEGER DEFAULT 0','Number of failed sync attempts'),
        ('last_error',      'TEXT',             'Last sync error message'),
        ('last_attempt',    'DATETIME',         'Last sync attempt timestamp'),
        ('ismanual',        'INTEGER DEFAULT 0','0 = face recognition, 1 = manual entry'),
        ('userimage',       'TEXT',             'Base64 snapshot taken at punch time'),
    ]
)

# --- FACEVECTOR_UPDATES ---
heading(doc, '4.3 facevector_updates – Cloud Sync Queue', 2)
section_table(doc,
    ['Column', 'Type', 'Description'],
    [
        ('id',          'INTEGER PK AUTOINCREMENT', 'Row ID'),
        ('uuid',        'TEXT NOT NULL',            'Employee GUID'),
        ('staffid',     'TEXT',                     'Employee staff ID'),
        ('vector',      'TEXT',                     'Average embedding to upload'),
        ('vectors',     'TEXT',                     'Multiple embeddings to upload'),
        ('img',         'TEXT',                     'Base64 photo to upload'),
        ('action',      'TEXT DEFAULT update',      '"create" or "update"'),
        ('created_at',  'DATETIME DEFAULT NOW',     'Enqueue timestamp'),
        ('sync_status', 'INTEGER DEFAULT 0',        '0 = pending, 1 = synced'),
        ('retry_count', 'INTEGER DEFAULT 0',        'Sync retry counter'),
        ('last_error',  'TEXT',                     'Last upload error'),
        ('last_attempt','DATETIME',                 'Last upload attempt'),
    ]
)

# --- USERPREFERENCES ---
heading(doc, '4.4 UserPreferences', 2)
section_table(doc,
    ['Column', 'Type', 'Description'],
    [
        ('id',                  'INTEGER PK DEFAULT 1', 'Singleton row'),
        ('colorPreference',     'TEXT',                 'UI theme colour selection'),
        ('languagePreference',  'TEXT',                 'Locale setting'),
    ]
)

heading(doc, '4.5 Database Indexes', 2)
section_table(doc,
    ['Index Name', 'Table', 'Columns', 'Purpose'],
    [
        ('idx_punch_uuid',              'punchrecord',         'uuid',             'Employee punch lookups'),
        ('idx_punch_uuid_date',         'punchrecord',         'uuid, punchdate',  'Daily punch queries'),
        ('idx_punch_project',           'punchrecord',         'projectid',        'Project-level reports'),
        ('idx_punch_sync',              'punchrecord',         'punchmode',        'Sync status filtering'),
        ('idx_facevector_uuid',         'facevector',          'uuid',             'Face template lookups'),
        ('idx_facevector_updates_sync', 'facevector_updates',  'sync_status',      'Pending upload queue'),
    ]
)

heading(doc, '4.6 Database Configuration', 2)
code_block(doc,
"""PRAGMA foreign_keys  = ON;      -- Enforce FK constraints
PRAGMA journal_mode  = WAL;     -- Allow concurrent reads during writes
PRAGMA synchronous   = NORMAL;  -- Balance durability vs. speed""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  5. CORE ALGORITHMS
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '5. Core Algorithms & Techniques', 1)
add_horizontal_rule(doc)

# 5.1
heading(doc, '5.1 Face Embedding Extraction (TensorFlow Lite)', 2)
body(doc,
    'A custom TensorFlow Lite model (TFLiteFaceModule) is invoked via a React Native native '
    'module. The pipeline processes each camera frame as follows:'
)
for step in [
    '1. Face detected in frame by ML Kit → bounding box returned.',
    '2. Image cropped to face region using @react-native-community/image-editor.',
    '3. Cropped image resized to model input dimensions via react-native-image-resizer.',
    '4. Pixel values normalised; image passed to TF Lite model.',
    '5. Output layer produces a 512-dimensional float32 embedding vector.',
    '6. Vector normalised to unit length (magnitude = 1) for cosine similarity.',
]:
    bullet(doc, step)

# 5.2
heading(doc, '5.2 Cosine Similarity (Face Matching)', 2)
body(doc,
    'All face comparisons use cosine similarity between 512-D vectors. '
    'A score of 1.0 means identical; 0.0 means completely unrelated.'
)
code_block(doc,
"""cosineSimilarity(a, b):
  dot   = Σ (a[i] × b[i])
  magA  = √(Σ a[i]²)
  magB  = √(Σ b[i]²)
  return dot / (magA × magB + ε)

Typical thresholds:
  ≥ 0.75  →  Very strong match (enrollment duplicate detection)
  ≥ 0.60  →  Match accepted (check-in/out recognition)
  ≥ 0.55  →  Loose match (configurable recognition threshold)
  < 0.55  →  No match → manual entry fallback""")

# 5.3
heading(doc, '5.3 Vector Normalisation', 2)
code_block(doc,
"""normalizeVector(v):
  sqSum  = Σ v[i]²
  invMag = 1 / √(sqSum + ε)
  return v.map(x → x × invMag)

After normalisation ||v|| = 1, so cosine similarity
reduces to the dot product — faster computation.""")

# 5.4
heading(doc, '5.4 Batched Linear Search', 2)
body(doc,
    'For datasets up to ~10,000 faces the app performs batched linear search to avoid '
    'blocking the JavaScript event loop:'
)
code_block(doc,
"""BATCH_SIZE = 3500 (configurable: 1000–5000)

For each batch of vectors:
  1. Compute cosine similarity against query embedding
  2. Track best (score, employee) tuple
  3. If best score > 0.70 → early exit (strong match found)
  4. yield to event loop (setImmediate / Promise.resolve)
  5. Continue with next batch

Result: top-1 nearest neighbour with confidence score""")

# 5.5
heading(doc, '5.5 HNSW – Approximate Nearest Neighbour (ANN) Index', 2)
body(doc,
    'The hnsw library provides an O(log n) alternative to linear search for very large '
    'datasets. Currently under development (code commented out).'
)
code_block(doc,
"""Configuration:
  dimension       = 512
  M               = 16    (connections per node)
  efConstruction  = 200   (build quality)
  efSearch        = 64    (query accuracy)

Search:
  index.search(queryVector, k=5)
  → returns k nearest IDs + distances""")

# 5.6
heading(doc, '5.6 Vector Blending – Continuous Learning', 2)
body(doc,
    'After each successful check-in the face model self-improves using exponential '
    'moving average blending:'
)
code_block(doc,
"""blendVectors(oldVector, newVector, α = 0.9):
  blended = oldVector.map((val, i) →
    α × val + (1 - α) × newVector[i]
  )
  return normalizeVector(blended)

α = 0.9  → 90% old model, 10% new observation
  Conservative update that prevents forgetting
  while adapting to gradual appearance changes.

Multiple embeddings (up to 5) also stored to
  represent different angles / lighting.""")

# 5.7
heading(doc, '5.7 Face Quality Scoring', 2)
body(doc,
    'Before capturing a face embedding the system scores the current camera frame '
    '(0–100 points) to ensure quality:'
)
section_table(doc,
    ['Factor', 'Weight', 'Description'],
    [
        ('Detection confidence', '20 pts', 'ML Kit face detection probability'),
        ('Head rotation (yaw/pitch/roll)', '30 pts', 'Penalises large head tilts'),
        ('Eye landmark alignment',  '25 pts', 'Horizontal symmetry of eye positions'),
        ('Face size ratio',         '15 pts', 'Face must occupy sufficient frame area'),
        ('Eye openness',            '10 pts', 'Detects closed / partially closed eyes'),
    ]
)
body(doc,
    'Capture locking: 3 consecutive high-quality frames required, followed by a '
    '300 ms stabilisation delay before the embedding is extracted.'
)

# 5.8
heading(doc, '5.8 Ray-Casting Point-in-Polygon (Geo-fencing)', 2)
body(doc,
    'Validates that the employee is physically within an assigned project boundary:'
)
code_block(doc,
"""isPointInPolygon(point, polygon, extraLeverage, minBuffer):
  1. Cast a horizontal ray from point to +∞
  2. Count how many polygon edges the ray crosses
  3. Odd count  → point is INSIDE polygon
  4. Even count → point is OUTSIDE polygon
  5. If outside: compute distance to nearest polygon edge
  6. If distance ≤ minBuffer (default 20 m) → still accept
     (tolerates GPS accuracy error)""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  6. MODULE-BY-MODULE BREAKDOWN
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '6. Module-by-Module Breakdown', 1)
add_horizontal_rule(doc)

# ─── face.service.js ───
heading(doc, '6.1 Face Service  (services/face.service.js)', 2)
body(doc, 'Central orchestrator for all face recognition operations.')
section_table(doc,
    ['Function', 'Description'],
    [
        ('loadVectorsService()',        'Loads all face vectors from SQLite into in-memory VECTOR_STORE'),
        ('recognizeFaceService(img)',   'Extracts embedding → batched search → returns best match + score'),
        ('enrollFaceService(data)',     'Validates duplicate → stores new vector in DB → queues cloud sync'),
        ('updateFaceService(data)',     'Replaces existing embedding; updates multi-vector array'),
        ('improveFaceModelService()',   'Post check-in learning: blends new embedding into stored template'),
    ]
)
body(doc, 'In-memory VECTOR_STORE structure:')
code_block(doc,
"""VECTOR_STORE = {
  data:        [...],  // All enrolled faces (complete DB)
  userData:    [...],  // Operator's own faces (priority match)
  lastUpdated: Date    // Version timestamp
}""")

# ─── faceProcessing.service.js ───
heading(doc, '6.2 Face Processing Service  (services/faceProcessing.service.js)', 2)
body(doc, 'Real-time frame evaluation; runs on every camera frame during check-in/enrollment.')
section_table(doc,
    ['Function', 'Description'],
    [
        ('evaluateFaceQuality(face, frame)', 'Scores frame 0–100; returns score + user feedback string'),
        ('resetFaceStabilizer()',            'Resets consecutive-frame counter to 0'),
    ]
)

# ─── attendance.service.js ───
heading(doc, '6.3 Attendance Service  (services/attendance.service.js)', 2)
section_table(doc,
    ['Function', 'Description'],
    [
        ('checkEmployeeTodayStatusService(uuid, projectId)', 'Returns today\'s punch status (in/out/none)'),
        ('checkActiveCheckinService(uuid)',                  'Checks if employee has an open shift'),
        ('processCheckInService(data)',                      'Validates + records a check-in punch'),
        ('processCheckOutService(data)',                     'Validates + records a check-out punch'),
        ('processCheckoutAndCheckinService(data)',           'Handles project-to-project transition'),
        ('manualEntryService(data)',                         'Records manual attendance (ismanual=1)'),
    ]
)

# ─── sync.service.js ───
heading(doc, '6.4 Sync Service  (services/sync.service.js)', 2)
section_table(doc,
    ['Function', 'Description'],
    [
        ('pullVectorsService(user, onProgress)', 'Paginates server API; inserts/updates local facevector rows'),
        ('pushVectorsService(user, onProgress)', 'Reads pending facevector_updates rows; uploads in batches of 50'),
    ]
)

# ─── backgroundSync.service.js ───
heading(doc, '6.5 Background Sync Service  (services/backgroundSync.service.js)', 2)
body(doc,
    'Uses react-native-background-fetch to schedule periodic sync tasks when the app is '
    'backgrounded. Monitors network connectivity and only syncs on available connection.'
)

# ─── dashboard.service.js ───
heading(doc, '6.6 Dashboard Service  (services/dashboard.service.js)', 2)
section_table(doc,
    ['Function', 'Description'],
    [
        ('getTodayStatsService(projectId)', 'Aggregates today\'s check-in/out counts from SQLite'),
        ('getOverallSyncStatsService()',    'Returns total synced vs. unsynced record counts'),
        ('checkInitialSyncService()',       'Determines if initial vector pull is required'),
    ]
)

# ─── employee.service.js ───
heading(doc, '6.7 Employee Service  (services/employee.service.js)', 2)
section_table(doc,
    ['Function', 'Description'],
    [
        ('searchEmployeeService(empId, token)', 'Queries cloud API for employee by staff ID'),
    ]
)

# ─── logger.service.js / bugfender.service.js ───
heading(doc, '6.8 Logging Services', 2)
body(doc,
    'logger.service.js captures session events to the local log.repository; '
    'bugfender.service.js forwards errors and crash data to the Bugfender cloud dashboard '
    'using @bugfender/rn-bugfender.'
)

# ─── Repositories ───
heading(doc, '6.9 Database Repositories', 2)
section_table(doc,
    ['Repository File', 'Key Methods'],
    [
        ('staff.repository.js',           'addStaff, getAllStaff, getStaffByCreatedBy, updateFaceVectors, getStaffImage'),
        ('punch.repository.js',           'recordPunch, checkTodayPunch, getActiveCheckIn, processCheckoutAndCheckin'),
        ('facevector_updates.repository.js', 'addFaceUpdate, getAllUnsynced, markSynced'),
        ('stats.repository.js',           'getTodayCheckinCount, getTodayCheckoutCount, getSyncStats'),
        ('log.repository.js',             'insertLog, getLogs'),
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  7. SCREEN / UI DOCUMENTATION
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '7. Screen / UI Documentation', 1)
add_horizontal_rule(doc)

screens = [
    {
        'num': '7.1',
        'name': 'Splash Screen',
        'file': 'screens/splash.js',
        'desc': (
            'First screen shown on app launch. Handles app initialisation, loads cached '
            'settings, checks stored authentication, and routes the user accordingly.'
        ),
        'features': [
            'Animated entrance: fade-in + scale animation',
            'Loads all configurable settings from AsyncStorage into memory',
            'Reads stored user session from AsyncStorage',
            'Routes to: Login (unauthenticated), Role Selection (multi-role user), Landing Page (single-role / cached role)',
            'Displays app version number',
        ]
    },
    {
        'num': '7.2',
        'name': 'Login Screen',
        'file': 'screens/login.js',
        'desc': (
            'Credential entry screen. Authenticates the user against the cloud API and '
            'persists the returned JWT token.'
        ),
        'features': [
            'Email / Employee ID text input',
            'Password input with show/hide toggle',
            'Async POST /api/login call',
            'Error alert on failure',
            'Stores user object (including JWT) to AsyncStorage on success',
            'Navigates to RoleSelection on success',
        ]
    },
    {
        'num': '7.3',
        'name': 'Role Selection Screen',
        'file': 'screens/roleSelection.js',
        'desc': (
            'Displayed when a user has multiple assigned roles. Allows them to choose '
            'which role to operate under for this session.'
        ),
        'features': [
            'Displays each role as a card with name, description, and icon',
            'Auto-selects and skips this screen if only one role exists',
            'Saves selected role to AsyncStorage',
            'Logout button clears all stored session data',
        ]
    },
    {
        'num': '7.4',
        'name': 'Landing Page  (Bottom Tab Navigator)',
        'file': 'screens/landingPage.js',
        'desc': (
            'Shell screen that hosts the five main tabs. Initialises continuous GPS tracking, '
            'geo-fencing validation, network monitoring, and background fetch.'
        ),
        'features': [
            'Five tabs: Dashboard, Check In, Check Out, Employees, History',
            'GPS continuous updates via LocationService',
            'Automatic project geo-fencing check on location change',
            'Network connectivity monitoring via @react-native-community/netinfo',
            'Background fetch initialisation for sync tasks',
            'Face vector pre-loading into VECTOR_STORE on mount',
        ]
    },
    {
        'num': '7.5',
        'name': 'Dashboard Screen',
        'file': 'screens/dashboard.js',
        'desc': 'Main overview screen showing real-time attendance statistics for the current day.',
        'features': [
            'Today\'s total check-ins and check-outs (from local DB)',
            'Synced vs. unsynced record counts',
            'Network status banner (online / offline)',
            'Last sync timestamp',
            'Staff profile card with photo',
            'Project dropdown selector',
            'Manual sync trigger button',
            'Pull-to-refresh',
        ]
    },
    {
        'num': '7.6',
        'name': 'Check-In Screen',
        'file': 'screens/checkin.js',
        'desc': (
            'Camera-based face recognition screen for recording employee arrivals. '
            'Guides the employee to position their face, captures 1–3 photos, runs '
            'recognition, and records the attendance punch.'
        ),
        'features': [
            'Camera permission request on mount',
            'Real-time face detection overlay (oval guide)',
            'Live quality feedback messages (e.g. "Move closer", "Look straight")',
            'Auto-capture when quality score is stable for 3 consecutive frames',
            'Configurable capture count (1–3 photos)',
            'Embedding extraction and batched vector search',
            'Single match → automatic attendance record + audio success sound',
            'Multiple matches → FaceConfirmationPopup for manual selection',
            'No match → ManualEntryModal fallback',
            'Front / rear camera toggle button',
            'ActiveCheckinPopup if employee already checked in to different project',
        ]
    },
    {
        'num': '7.7',
        'name': 'Check-Out Screen',
        'file': 'screens/checkOut.js',
        'desc': (
            'Identical workflow to Check-In but records a "out" punch. '
            'Validates that the employee has an active (open) check-in before allowing checkout.'
        ),
        'features': [
            'Same face recognition pipeline as Check-In',
            'Verifies active check-in via checkActiveCheckinService',
            'Warning alert if no active check-in found',
        ]
    },
    {
        'num': '7.8',
        'name': 'Face Enrollment Screen',
        'file': 'screens/FaceEnrollmentScreen.js',
        'desc': (
            'Multi-step guided face enrollment for new employees. '
            'Collects multiple face angles to build a robust face template.'
        ),
        'features': [
            'Step-by-step guidance: "Look straight", "Turn left", "Turn right", etc.',
            'Configurable capture count (3 or 5 angles)',
            'Per-step quality feedback',
            'Progress indicator (step N of total)',
            'Manual snap button for each step',
            'Duplicate face detection before saving (threshold: 0.75)',
            'Averages all captured embeddings into a single template vector',
            'Stores both average vector and all individual vectors',
            'Queues new enrollment for cloud sync',
            'Image compression before storage',
        ]
    },
    {
        'num': '7.9',
        'name': 'Employee Management Screen',
        'file': 'screens/employee.js',
        'desc': (
            'Two-tab screen for managing the employee face database.'
        ),
        'features': [
            'Tab 1 – Search & Enroll: search employee by staff ID, fetch from cloud API, initiate enrollment',
            'Tab 2 – Enrolled Employees: paginated list of all enrolled faces, search filter, view details, update template',
            'EmployeeDetailsModal: shows employee info + enrolled face photo',
            'Delete enrollment option (admin role)',
        ]
    },
    {
        'num': '7.10',
        'name': 'Attendance History Screen',
        'file': 'screens/histroy.js',
        'desc': 'Paginated log of all attendance records stored locally.',
        'features': [
            'Date filter (calendar picker)',
            'Employee filter',
            'Sync status filter (synced / unsynced)',
            'Shows: employee name, check-in/out times, project, GPS coordinates',
            'Manual vs. face recognition indicator badge',
            'Pull-to-refresh',
            'Pagination for large datasets',
        ]
    },
    {
        'num': '7.11',
        'name': 'Sync Data Screen',
        'file': 'screens/syncdata.js',
        'desc': 'Manual synchronisation control panel.',
        'features': [
            'Network status indicator (must be online to sync)',
            'Pull Vectors button – downloads all server face templates',
            'Push Vectors button – uploads pending local face updates',
            'Progress bar with record count',
            'Navigation locked during active sync (prevents interruption)',
            'Success / error alert on completion',
        ]
    },
    {
        'num': '7.12',
        'name': 'Settings Screen',
        'file': 'screens/SettingsScreen.js',
        'desc': 'In-app configuration panel for tuning recognition parameters.',
        'features': [
            'Slider: Face Match Threshold (0.1–1.0, step 0.05)',
            'Slider: Recognition Threshold (0.1–1.0)',
            'Slider: Enrollment Duplicate Threshold (0.1–1.0)',
            'Slider: Template Update Threshold (0.1–1.0)',
            'Dropdown: Enrollment Capture Count (3 or 5)',
            'Dropdown: Verification Capture Count (1, 2, or 3)',
            'Dropdown: Batch Size (1000, 2000, 3500, 5000)',
            'Save persists to AsyncStorage; Cancel reverts changes',
            'Helper description text for each setting',
        ]
    },
]

for s in screens:
    heading(doc, f"{s['num']} {s['name']}  ({s['file']})", 2)
    body(doc, s['desc'])
    for f in s['features']:
        bullet(doc, f)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  8. NAVIGATION & APPLICATION FLOW
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '8. Navigation & Application Flow', 1)
add_horizontal_rule(doc)

heading(doc, '8.1 Navigation Stack Overview', 2)
code_block(doc,
"""Root Stack Navigator
├── Splash                   ← Always first; decides routing
├── Login
├── RoleSelection
└── LandingPage              ← Bottom Tab Navigator
    ├── Dashboard
    ├── CheckIn
    ├── CheckOut
    ├── Employees
    │   ├── EnrollEmployeeTab   (Search & Enroll)
    │   └── EnrolledEmployeeTab (Enrolled List)
    └── History

Modal Screens (push on top of tabs):
    ├── FaceEnrollmentScreen
    ├── SettingsScreen
    └── SyncData""")

heading(doc, '8.2 Authentication Flow', 2)
for step in [
    '1. App starts → Splash screen mounts.',
    '2. Splash loads cached settings from AsyncStorage.',
    '3. Splash checks for stored user object in AsyncStorage.',
    '4. If no user → navigate to Login.',
    '5. User enters credentials → POST /api/login.',
    '6. On success → save user + token to AuthContext + AsyncStorage.',
    '7. If user has multiple roles → navigate to RoleSelection.',
    '8. User selects role → save to AsyncStorage → navigate to LandingPage.',
    '9. On subsequent launches steps 4–8 are skipped (cached session).',
]:
    bullet(doc, step)

heading(doc, '8.3 Check-In Flow (Detailed)', 2)
for step in [
    '1.  LandingPage initialises GPS and loads face vectors into VECTOR_STORE.',
    '2.  User taps Check In tab.',
    '3.  Camera permission requested if not granted.',
    '4.  VisionCamera preview starts; ML Kit face detection begins on each frame.',
    '5.  evaluateFaceQuality() scores each frame (0–100).',
    '6.  Real-time feedback shown ("Move closer", "Hold still", etc.).',
    '7.  Three consecutive high-quality frames detected → 300 ms stabilisation wait.',
    '8.  Camera captures 1–3 photos (configurable CAPTURE_COUNT).',
    '9.  TFLiteFaceModule extracts 512-D embedding from each photo.',
    '10. Embeddings averaged; resulting vector normalised.',
    '11. Batched cosine similarity search across VECTOR_STORE.',
    '12a. Single best match (score ≥ threshold) → processCheckInService() called.',
    '12b. Multiple candidates → FaceConfirmationPopup shown for user to select.',
    '12c. No match → ManualEntryModal opens for employee ID entry.',
    '13. punchrecord row inserted into SQLite (punchtype = "in").',
    '14. facevector_updates row created → queued for cloud push.',
    '15. improveFaceModelService() blends new embedding (continuous learning).',
    '16. Success audio played (react-native-sound-player).',
    '17. User navigated back to Dashboard.',
]:
    bullet(doc, step)

heading(doc, '8.4 Face Enrollment Flow', 2)
for step in [
    '1.  Admin searches for employee by staff ID on Employees tab.',
    '2.  Employee fetched from cloud API; details shown in modal.',
    '3.  Admin taps "Enroll Face" → FaceEnrollmentScreen opens.',
    '4.  Screen guides through N steps (3 or 5, configurable).',
    '5.  For each step: quality feedback shown; admin taps snap button.',
    '6.  Embedding extracted from each captured photo.',
    '7.  After all steps: duplicate check against VECTOR_STORE (threshold 0.75).',
    '8.  If duplicate found → error shown, enrollment blocked.',
    '9.  All embeddings averaged → stored as primary vector.',
    '10. Individual embeddings stored in vectors array (up to 5).',
    '11. Face photo compressed → stored as base64 in img field.',
    '12. Row inserted/updated in facevector table.',
    '13. Row inserted into facevector_updates with action = "create".',
    '14. VECTOR_STORE updated in memory.',
]:
    bullet(doc, step)

heading(doc, '8.5 Sync Flow', 2)
body(doc, 'Pull (Download):')
for step in [
    '1. GET server vector count.',
    '2. Paginate: POST /api/getallvectors with page=1,2,3…',
    '3. Each page: upsert rows in facevector SQLite table.',
    '4. Update VECTOR_STORE in memory.',
    '5. Save sync timestamp to AsyncStorage.',
]:
    bullet(doc, step)
body(doc, 'Push (Upload):')
for step in [
    '1. Query facevector_updates WHERE sync_status = 0.',
    '2. Batch 50 records → POST /api/multipleSaveentrolledimage.',
    '3. On success: UPDATE sync_status = 1.',
    '4. On failure: increment retry_count; save last_error.',
]:
    bullet(doc, step)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  9. API INTEGRATION
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '9. API Integration', 1)
add_horizontal_rule(doc)

body(doc, 'Base URL:  http://44.204.69.221  (configured in config/config.ts)')
body(doc, 'All authenticated endpoints require: Authorization: Bearer <jwt_token>')
doc.add_paragraph()

apis = [
    {
        'endpoint': 'POST /api/login',
        'purpose':  'Authenticate user; returns JWT token',
        'request':  '{ email, password }',
        'response': '{ data: { data: { ...user }, access_token: "jwt" } }',
    },
    {
        'endpoint': 'POST /api/employees',
        'purpose':  'Search employee by staff ID',
        'request':  '{ emp_id }',
        'response': '{ data: [ { employee objects } ] }',
    },
    {
        'endpoint': 'POST /api/getallvectors',
        'purpose':  'Paginated download of face templates',
        'request':  '{ update_date, createdby, length: 50, page: N }',
        'response': '{ data: { data: [ { empguid, vector, vectors, image, ... } ], total_count } }',
    },
    {
        'endpoint': 'POST /api/multipleSaveentrolledimage',
        'purpose':  'Upload batch of enrolled face updates',
        'request':  '{ enrollments: [ { uuid, staffid, vector, vectors, image } ] }',
        'response': '{ success: true }',
    },
    {
        'endpoint': 'POST /api/savepunchrecords',
        'purpose':  'Upload attendance punch records',
        'request':  '[ { guid, emp_id, project_id, attendance_type, date, checkin_time, checkout_time, checkin_lat, checkin_lang, checkout_lat, checkout_lang, checkin_image, checkout_image, ismanual } ]',
        'response': '{ success: true }',
    },
    {
        'endpoint': 'POST /api/history',
        'purpose':  'Fetch attendance history from server',
        'request':  '{ filters, page, limit }',
        'response': '{ data: [ punch records ] }',
    },
]

for api in apis:
    heading(doc, api['endpoint'], 3)
    info_table(doc, [
        ('Purpose',   api['purpose']),
        ('Request',   api['request']),
        ('Response',  api['response']),
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  10. AUTHENTICATION & SECURITY
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '10. Authentication & Security', 1)
add_horizontal_rule(doc)

heading(doc, '10.1 Authentication Method', 2)
section_table(doc,
    ['Aspect', 'Detail'],
    [
        ('Protocol',           'JWT (JSON Web Token) over HTTP REST'),
        ('Credentials',        'Email + Password'),
        ('Token Storage',      'AsyncStorage (local device; not in secure keychain)'),
        ('Session Persistence','Token re-loaded from AsyncStorage on each app launch'),
        ('Logout',             'AsyncStorage cleared; AuthContext reset; navigate to Login'),
    ]
)

heading(doc, '10.2 Role-Based Access Control', 2)
body(doc,
    'Each user account has one or more roles returned from the login API. Each role object '
    'contains:'
)
for item in [
    'name – display name of the role (e.g. "Attendance Admin")',
    'permissions – array of allowed actions',
    'attendance_logic – rules such as project_required, allowed_punch_types',
    'The selected role is stored in AsyncStorage and governs what screens and actions are accessible.',
]:
    bullet(doc, item)

heading(doc, '10.3 Biometric Security Layers', 2)
section_table(doc,
    ['Layer', 'Mechanism'],
    [
        ('Face Recognition',    'TF Lite 512-D embedding + cosine similarity (threshold configurable)'),
        ('Duplicate Prevention','Enrollment blocked if similarity ≥ 0.75 to any existing face'),
        ('Anti-Spoofing',       'FaceSpoofDetector native Android module (integrated but not yet active)'),
        ('Geo-fencing',         'Ray-casting polygon check; must be within project boundary to punch'),
        ('Manual Fallback',     'Marked ismanual=1; distinguishable in reports'),
        ('Liveness Quality',    'Multi-factor frame scoring prevents low-quality / blur captures'),
    ]
)

heading(doc, '10.4 Known Security Considerations', 2)
for item in [
    'HTTP (not HTTPS): Current Base_URL uses plain HTTP; should be upgraded for production.',
    'Token in AsyncStorage: Not stored in Android Keystore / iOS Secure Enclave.',
    'Anti-spoofing: Module exists but is not yet wired into the main recognition flow.',
    'Image storage: Face photos stored as base64 in SQLite; encryption at rest not implemented.',
]:
    bullet(doc, item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  11. STATE MANAGEMENT
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '11. State Management', 1)
add_horizontal_rule(doc)

section_table(doc,
    ['Storage', 'Used For', 'Scope'],
    [
        ('React Context (AuthContext)',     'Current user object, login/logout actions',           'App-wide; cleared on restart'),
        ('AsyncStorage',                   'User session, selected role, project, sync date, settings', 'Persisted across restarts'),
        ('In-Memory VECTOR_STORE',         'All face embeddings for fast recognition lookup',     'App session; reloaded on launch'),
        ('SQLite',                         'All operational data: punches, faces, sync queue',    'Permanent local storage'),
        ('React component state (useState)', 'UI state: form values, loading flags, modal visibility', 'Component lifetime'),
    ]
)

body(doc, 'No Redux, MobX, or Zustand is used. State is kept close to where it is needed.')

# ════════════════════════════════════════════════════════════════════════════════
#  12. SETTINGS & CONFIGURATION
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '12. Settings & Configuration', 1)
add_horizontal_rule(doc)

heading(doc, '12.1 Configurable Recognition Parameters', 2)
section_table(doc,
    ['Setting Key', 'Default', 'Range', 'Description'],
    [
        ('FACE_MATCH_THRESHOLD',           '0.60', '0.1–1.0',  'Minimum cosine similarity to accept a match'),
        ('RECOGNITION_THRESHOLD',          '0.55', '0.1–1.0',  'Looser threshold for daily check-in/out'),
        ('ENROLLMENT_DUPLICATE_THRESHOLD', '0.75', '0.1–1.0',  'Similarity above which enrollment is blocked'),
        ('TEMPLATE_UPDATE_THRESHOLD',      '0.60', '0.1–1.0',  'Minimum similarity to trigger learning update'),
        ('ENROLLMENT_CAPTURE_COUNT',       '3',    '3 or 5',   'Number of face angles captured during enrollment'),
        ('CAPTURE_COUNT',                  '2',    '1, 2, 3',  'Photos taken per check-in/out attempt'),
        ('BATCH_SIZE',                     '3500', '1000–5000','Vectors processed per batch during search'),
    ]
)

heading(doc, '12.2 App Configuration Files', 2)
section_table(doc,
    ['File', 'Purpose'],
    [
        ('config/config.ts',      'API Base_URL definition'),
        ('constants/colors.ts',   'Colour palette tokens (#0a2342, #00b8a9, etc.)'),
        ('app.json',              'Expo app name, icon, splash screen config'),
        ('babel.config.js',       'Babel transpilation presets'),
        ('metro.config.js',       'Metro bundler configuration'),
        ('tsconfig.json',         'TypeScript compiler options (optional typing)'),
        ('.eslintrc.js',          'ESLint rules'),
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  13. PERFORMANCE OPTIMISATIONS
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '13. Performance Optimisations', 1)
add_horizontal_rule(doc)

perf_items = [
    ('In-Memory Vector Cache',
     'VECTOR_STORE holds all face embeddings in RAM, eliminating per-recognition DB reads. '
     'Refreshed only after enrollment or sync.'),
    ('Batched Search with Event-Loop Yield',
     'Linear vector search runs in configurable batches (default 3500). Between batches, '
     'setImmediate / Promise.resolve yields control to prevent UI freezing.'),
    ('Early Exit Optimisation',
     'If a match score exceeds 0.70 during batched search, the loop exits immediately '
     'without scanning remaining batches.'),
    ('User-First Matching',
     'The current operator\'s own enrolled faces are checked first before the full database, '
     'allowing faster self-identification in use-cases where operators also punch.'),
    ('SQLite WAL Mode',
     'Write-Ahead Logging enables concurrent reads during writes, improving responsiveness '
     'during sync operations.'),
    ('Indexed Queries',
     'Six targeted indexes on punchrecord and facevector tables minimise full-table scans '
     'for common query patterns.'),
    ('Image Compression',
     'Face photos are compressed before base64 encoding to reduce database size and '
     'network payload during sync.'),
    ('Settings Cache',
     'All configurable settings are loaded into a module-level JS object at startup; '
     'subsequent reads are O(1) without AsyncStorage I/O.'),
    ('Lazy Image Loading',
     'LazyImage component defers image decoding until scroll position requires it, '
     'keeping list scroll smooth.'),
    ('React Native Reanimated',
     'Animations (splash, progress) run on the native UI thread via Reanimated 3, '
     'achieving 60 fps without blocking the JS thread.'),
    ('Capture Stabilisation',
     'Three consecutive high-quality frames required before capture, preventing wasted '
     'embedding extractions on blurry or partially visible faces.'),
]

for title, desc in perf_items:
    body(doc, title, bold=True)
    body(doc, desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
#  14. BACKGROUND SERVICES
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '14. Background Services', 1)
add_horizontal_rule(doc)

heading(doc, '14.1 Background Fetch (react-native-background-fetch)', 2)
for item in [
    'Registers a periodic task that executes even when the app is backgrounded.',
    'Checks network connectivity via NetInfo.',
    'Triggers pushVectorsService() to upload pending face updates.',
    'Triggers punch record sync to upload unsynced attendance.',
    'Respects platform constraints (Android Doze / iOS background execution limits).',
]:
    bullet(doc, item)

heading(doc, '14.2 Background Sync Service (services/backgroundSync.service.js)', 2)
for item in [
    'Listens to network state changes.',
    'On transition from offline → online: initiates sync queue flush.',
    'Exponential backoff on repeated failures.',
    'Logs sync results via logger.service.',
]:
    bullet(doc, item)

heading(doc, '14.3 Location Service (utils/LocationService.js)', 2)
for item in [
    'Wraps native Android LocationModule.',
    'Provides subscribe() / unsubscribe() event emitter API.',
    'Delivers: { latitude, longitude, accuracy } objects to subscribers.',
    'Used by LandingPage to continuously update current position for geo-fencing.',
]:
    bullet(doc, item)

# ════════════════════════════════════════════════════════════════════════════════
#  15. KNOWN ISSUES & LIMITATIONS
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '15. Known Issues & Limitations', 1)
add_horizontal_rule(doc)

section_table(doc,
    ['Issue', 'Detail', 'Status'],
    [
        ('HNSW index disabled',
         'hnsw library is integrated but the index build and search code is commented out. '
         'Falls back to O(n) linear search; impacts performance for > 10,000 enrolled faces.',
         'In development'),
        ('HTTP instead of HTTPS',
         'Base_URL uses plain HTTP; data transmitted in clear text. Vulnerable to MITM.',
         'Should be fixed before production'),
        ('GPS column named "lan"',
         'punchrecord schema uses column "lan" for longitude instead of conventional "lng" or "lon". '
         'Schema is stable; changing would require migration.',
         'Established naming'),
        ('Anti-spoofing not active',
         'FaceSpoofDetector native module is compiled and available but not called in the '
         'main recognition flow.',
         'In development'),
        ('JWT in AsyncStorage',
         'Auth token is stored in AsyncStorage (plaintext on device) rather than the Android '
         'Keystore or iOS Secure Enclave.',
         'Security improvement needed'),
        ('iOS feature parity',
         'FaceSpoofDetector and some native location features are Android-only. TSFModuleIOS '
         'provides partial iOS support.',
         'Platform limitation'),
        ('Background location',
         'Continuous GPS tracking works in foreground; background location beyond app '
         'background-fetch windows is partially limited.',
         'In development'),
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
#  16. BUILD & DEPLOYMENT
# ════════════════════════════════════════════════════════════════════════════════
heading(doc, '16. Build, Run & Deployment', 1)
add_horizontal_rule(doc)

heading(doc, '16.1 Prerequisites', 2)
for item in [
    'Node.js >= 18',
    'JDK 17 (for Android build)',
    'Android Studio + Android SDK 34',
    'Xcode 15+ (for iOS build)',
    'React Native CLI',
    'Expo CLI (for managed features)',
]:
    bullet(doc, item)

heading(doc, '16.2 Development Commands', 2)
code_block(doc,
"""# Install dependencies
npm install

# Start Metro bundler
npm start

# Build & run on Android emulator / device
npm run android   # = react-native run-android

# Build & run on iOS simulator
npm run ios       # = react-native run-ios

# Lint
npm run lint      # ESLint

# Test
npm test          # Jest""")

heading(doc, '16.3 Release Build (Android)', 2)
code_block(doc,
"""cd android
./gradlew assembleRelease          # APK
./gradlew bundleRelease            # AAB (Play Store)

Output: android/app/build/outputs/""")

heading(doc, '16.4 Native Modules Required', 2)
body(doc,
    'The following native modules must be compiled with the Android build. '
    'They are not available as standalone npm packages:'
)
section_table(doc,
    ['Module', 'Language', 'Purpose'],
    [
        ('TFLiteFaceModule',  'Java/Kotlin', '512-D face embedding extraction via TF Lite'),
        ('TFLiteModule',      'Java/Kotlin', 'General TF Lite inference'),
        ('TSFModuleIOS',      'Swift/ObjC',  'iOS TF Lite wrapper'),
        ('FaceSpoofDetector', 'Java/Kotlin', 'Anti-spoofing liveness detection (Android)'),
        ('LocationModule',    'Java/Kotlin', 'Native GPS location updates'),
    ]
)

# ─── Final page ─────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, 'Document Information', 1)
add_horizontal_rule(doc)
info_table(doc, [
    ('Document Title',  'FRAS – Face Recognition Attendance System – Technical Documentation'),
    ('Version',         '1.0'),
    ('App Version',     '1.1.4'),
    ('Generated On',    datetime.date.today().strftime('%B %d, %Y')),
    ('Prepared By',     'Claude Code (Anthropic)'),
    ('Framework',       'React Native 0.80.0 / Expo 53 / React 19.1.0'),
    ('Platform',        'Android (primary), iOS (supported)'),
])

# ── Save ─────────────────────────────────────────────────────────────────────────
out_path = r'd:\Fras Enchanced\FRAS-dev\FRAS_Documentation.docx'
doc.save(out_path)
print(f"Document saved: {out_path}")
